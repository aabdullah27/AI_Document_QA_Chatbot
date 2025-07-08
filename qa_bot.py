import streamlit as st
import faiss
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.groq import Groq
from docx import Document as DocxDocument
import pymupdf4llm
import tempfile
import os

# Initialize session state
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "index" not in st.session_state:
    st.session_state.index = None

def read_file(file):
    if file.name.endswith('.pdf'):
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_file_path = tmp_file.name

        try:
            # Use pymupdf4llm with the temporary file path
            md_text = pymupdf4llm.to_markdown(tmp_file_path)
        finally:
            # Clean up the temporary file
            os.unlink(tmp_file_path)

        return md_text
    elif file.name.endswith('.docx'):
        return "\n".join(para.text for para in DocxDocument(file).paragraphs)
    else:
        return file.getvalue().decode()

def process_documents(uploaded_files):
    documents = [Document(text=read_file(file), metadata={"filename": file.name}) for file in uploaded_files]
    
    d = 384  # Dimension of the embedding model
    faiss_index = faiss.IndexFlatL2(d)
    vector_store = FaissVectorStore(faiss_index=faiss_index)
    
    st.session_state.index = VectorStoreIndex.from_documents(documents, vector_store=vector_store)

def get_bot_response(user_input):
    if st.session_state.index is None:
        return "Please upload some documents first!"
    
    query_engine = st.session_state.index.as_query_engine(
        response_mode="compact"
    )
    response = query_engine.query(user_input)
    return str(response)

def clear_chat_history():
    st.session_state.chat_history = []

# Streamlit UI setup
st.set_page_config(page_title="Document QA Bot", page_icon="📚", layout="wide")

# Sidebar for API key input and document upload
with st.sidebar:
    st.header("🔑 API Key")
    st.session_state.api_key = st.text_input("Enter your Groq API Key", type="password")

    st.header("📁 Document Upload")
    uploaded_files = st.file_uploader(
        "Upload your documents (PDF, DOCX, TXT)",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'txt']
    )
    
    if uploaded_files and st.session_state.api_key:
        with st.spinner("Processing documents..."):
            # Initialize Groq LLM and embedding model only if API key is provided
            llm = Groq(api_key=st.session_state.api_key, model="llama-3.3-70b-versatile")
            embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
            
            # Configure global settings
            Settings.embed_model = embed_model
            Settings.llm = llm
            
            process_documents(uploaded_files)
        st.success(f"{len(uploaded_files)} document(s) processed successfully!")
    
    st.header("🧹 Clear Chat")
    if st.button("Clear Chat History"):
        clear_chat_history()
        st.rerun()

# Main content area
st.title("📚 Document QA Bot")
st.write("Ask questions about your uploaded documents!")

# Chat interface
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

user_input = st.chat_input("Ask a question about your documents...")

if user_input:
    st.session_state.chat_history.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)
    
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        full_response = get_bot_response(user_input)
        message_placeholder.markdown(full_response)
    
    st.session_state.chat_history.append({"role": "assistant", "content": full_response})

if not st.session_state.chat_history:
    st.info("Upload documents in the sidebar, then ask questions to get started!")
